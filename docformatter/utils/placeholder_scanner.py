"""
占位符扫描器
扫描 Word 文档中的 {{占位符}} 模式
"""

import re
from typing import List, Set
from docx import Document

PLACEHOLDER_PATTERN = re.compile(r'\{\{(.+?)\}\}')


def scan_placeholders(docx_path: str) -> List[str]:
    """扫描 docx 中的所有 {{占位符}}，返回去重排序的占位符名列表"""
    doc = Document(docx_path)
    found: Set[str] = set()

    # 扫描正文段落
    for para in doc.paragraphs:
        _scan_runs(para.runs, found)

    # 扫描表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_runs(para.runs, found)

    # 扫描页眉页脚
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan_runs(para.runs, found)
        for para in section.footer.paragraphs:
            _scan_runs(para.runs, found)

    return sorted(found)


def _scan_runs(runs, found: Set[str]):
    """扫描一组 run 中的占位符"""
    # 先合并同一段落的所有 run 文本，避免占位符被拆散
    full_text = ''.join(r.text for r in runs)
    for m in PLACEHOLDER_PATTERN.finditer(full_text):
        found.add(m.group(1))
