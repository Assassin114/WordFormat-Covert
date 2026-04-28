"""
表格格式化处理器
支持：常规表格、测试记录表、试验记录表
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..utils.logger import get_logger

logger = get_logger()

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class TableHandler:
    """表格格式化处理器"""
    
    def __init__(self, template_config_path: str = None):
        self.template_config_path = template_config_path
        self.load_template()
    
    def load_template(self):
        """加载字段映射模板"""
        if self.template_config_path is None:
            template_file = Path(__file__).parent.parent / "templates" / "table_field_template.json"
        else:
            template_file = Path(self.template_config_path)
        
        if template_file.exists():
            with open(template_file, 'r', encoding='utf-8') as f:
                self.template = json.load(f)
        else:
            self.template = self._get_default_template()
    
    def _get_default_template(self):
        """获取默认模板"""
        return {
            "tables": {
                "常规表格": {
                    "header_style": {"bold": True, "align": "center"},
                    "body_style": {"align_rule": "text_length", "short_text_align": "center", "long_text_align": "left", "threshold_chars": 10}
                },
                "测试记录表": {
                    "section_header": {"fields": ["基本信息", "测试过程", "测试结论"], "bold": True, "align": "center"},
                    "field_name": {"fields": ["测试用例名称", "用例测试用例标识", "测试内容", "测试方法", "前提和约束", "终止条件"], "bold": True, "align": "left"},
                    "field_value": {"bold": False, "align": "left"},
                    "column_header": {"fields": ["序号", "操作步骤", "预期结果", "实测结果", "是否通过"], "bold": True, "align": "center"},
                    "data_row": {"bold": False, "align": "left"},
                    "signature_row": {"fields": ["测试人员", "操作人员", "测试时间"], "name_bold": True, "name_align": "center", "value_bold": False, "value_align": "left"}
                }
            }
        }
    
    def detect_table_type(self, table) -> str:
        """检测表格类型"""
        text = self._get_table_text(table)
        
        record_indicators = ["基本信息", "测试用例", "测试内容", "测试方法", "前提和约束",
                            "终止条件", "测试过程", "试验科目", "试验目的", "试验流程",
                            "测试人员", "操作人员", "测试时间"]
        for indicator in record_indicators:
            if indicator in text:
                return "测试记录表"
        
        trial_indicators = ["试验标识", "试验结论", "试验人员", "完成时间"]
        for indicator in trial_indicators:
            if indicator in text:
                return "试验记录表"
        
        return "常规表格"
    
    def _get_table_text(self, table) -> str:
        """获取表格所有文本"""
        texts = []
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            texts.append(run.text.strip())
        return " ".join(texts)
    
    def format_table(self, table, table_type: str = None):
        """格式化表格"""
        if table_type is None:
            table_type = self.detect_table_type(table)
        
        if table_type in ["测试记录表", "试验记录表"]:
            self._format_record_table(table, table_type)
        else:
            self._format_regular_table(table)
    
    def _format_regular_table(self, table):
        """格式化常规表格"""
        cfg = self.template["tables"]["常规表格"]
        
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    if row_idx == 0:  # 表头
                        self._apply_style(para, cfg["header_style"])
                    else:  # 正文
                        threshold = cfg["body_style"]["threshold_chars"]
                        align = cfg["body_style"]["short_text_align"] if len(text) <= threshold else cfg["body_style"]["long_text_align"]
                        style = cfg["body_style"].copy()
                        style["align"] = align
                        self._apply_style(para, style)
    
    def _format_record_table(self, table, table_type: str):
        """格式化记录表"""
        cfg = self.template["tables"].get(table_type, self.template["tables"]["测试记录表"])
        
        # 收集所有已知字段名
        section_fields = cfg.get("section_header", {}).get("fields", [])
        field_name_fields = cfg.get("field_name", {}).get("fields", [])
        sig_fields = cfg.get("signature_row", {}).get("fields", [])
        col_fields = cfg.get("column_header", {}).get("fields", [])
        known_fields = set(section_fields + field_name_fields + sig_fields + col_fields)
        
        name_style = {k: v for k, v in cfg.get("field_name", {}).items() if k != "fields"}
        name_style["bold"] = True
        value_style = {k: v for k, v in cfg.get("field_value", {}).items()}
        data_style = cfg.get("data_row", {})
        
        for row_idx, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            row_text_combined = " ".join(row_texts)
            
            if not row_text_combined:
                continue
            
            # 分类行类型
            is_section = any(f in row_text_combined for f in section_fields)
            is_col_header = row_texts and row_texts[0] in col_fields
            is_sig = any(f in row_texts[:3] for text in row_texts[:3] for f in sig_fields if f in text)
            is_data = row_texts and row_texts[0].isdigit()
            
            if is_section:
                self._format_section_header_row(row, cfg)
            elif is_col_header:
                self._format_column_header_row(row, name_style)
            elif is_sig:
                self._format_signature_row(row, cfg, sig_fields, name_style, value_style)
            elif is_data:
                self._format_data_row(row, data_style)
            else:
                # 字段行
                self._format_field_row(row, known_fields, name_style, value_style)
    
    def _format_section_header_row(self, row, cfg):
        style = cfg["section_header"]
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    self._apply_style(para, style)
    
    def _format_column_header_row(self, row, name_style):
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    self._apply_style(para, name_style)
    
    def _format_field_row(self, row, known_fields, name_style, value_style):
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            is_name = cell_text in known_fields
            for para in cell.paragraphs:
                self._apply_style(para, name_style if is_name else value_style)
    
    def _format_signature_row(self, row, cfg, sig_fields, name_style, value_style):
        sig_set = set(sig_fields)
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            is_name = cell_text in sig_set or any(f in cell_text for f in sig_fields)
            for para in cell.paragraphs:
                self._apply_style(para, name_style if is_name else value_style)
    
    def _format_data_row(self, row, data_style):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = data_style.copy()
                if cell_idx == 0 and text.isdigit():
                    style["align"] = "center"
                self._apply_style(para, style)
    
    def _apply_style(self, para, style):
        """应用样式到段落"""
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "both": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        if "align" in style and style["align"] in align_map:
            para.alignment = align_map[style["align"]]
        
        for run in para.runs:
            if "bold" in style:
                run.bold = style["bold"]
            if "font" in style:
                run.font.name = style["font"]
                run._element.rPr.rFonts.set(qn('w:eastAsia'), style["font"])
            if "size" in style:
                run.font.size = Pt(style["size"] / 2)
