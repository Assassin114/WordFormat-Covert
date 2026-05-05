"""
Word→Markdown 转换器 (Phase B-1)
结构化提取 Word 文档：YAML frontmatter + section 标记 + 书签 + REF域 + 脚注 + Visio
"""

import re
import os
import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from ..utils.logger import get_logger
from .document_analyzer import DocumentAnalyzer, SectionType, ElementType, ImageKind, EquationKind

logger = get_logger()

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
OLE_NS = '{urn:schemas-microsoft-com:office:office}'
REL_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
DML_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
WPD_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
MC_NS = '{http://schemas.openxmlformats.org/markup-compatibility/2006}'


class Word2MDConverter:
    """Word→Markdown 结构化转换器"""

    def __init__(self):
        self.analyzer = DocumentAnalyzer()
        self.image_counter = 0
        self.table_counter = 0
        self._bookmarks: dict = {}
        self._ref_fields: list = []
        self._footnotes: dict = {}
        self._footnote_counter = 0

    def convert(self, docx_path: str, output_dir: str) -> str:
        """
        将 Word 文档转换为 L3 级 Markdown（包含所有内部信息）

        Args:
            docx_path: 输入 .docx 路径
            output_dir: 输出目录（存放 .md、images/、attachments/）

        Returns:
            输出 .md 文件路径
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        images_dir = output_dir / "images"
        images_dir.mkdir(exist_ok=True)
        attachments_dir = output_dir / "attachments"
        attachments_dir.mkdir(exist_ok=True)

        self.image_counter = 0
        self.table_counter = 0
        self._bookmarks = {}
        self._ref_fields = []
        self._footnotes = {}
        self._footnote_counter = 0

        doc = Document(docx_path)
        structure = self.analyzer.analyze(docx_path)

        self._scan_bookmarks(doc)
        self._scan_ref_fields(doc)
        self._scan_footnotes(doc)

        md_lines = []
        md_lines.append("---")
        md_lines.extend(self._build_frontmatter(doc, structure))
        md_lines.append("---")
        md_lines.append("")

        for section in structure.sections:
            self._convert_section(section, doc, md_lines, images_dir, attachments_dir, docx_path)

        # 统一处理表格（全局一次，避免 section 重复）
        self._convert_all_tables(doc, md_lines)

        # 添加脚注定义
        self._append_footnote_definitions(md_lines)

        md_content = "\n".join(md_lines)
        md_name = Path(docx_path).stem + ".md"
        md_path = output_dir / md_name
        md_path.write_text(md_content, encoding="utf-8")

        logger.info(f"Word→MD 完成: {md_path}")
        return str(md_path)

    # ==================== Frontmatter ====================

    def _build_frontmatter(self, doc, structure) -> list:
        lines = []
        has_cover = any(s.section_type == SectionType.COVER for s in structure.sections)
        has_signature = any(s.section_type == SectionType.SIGNATURE for s in structure.sections)
        has_revision = any(s.section_type == SectionType.REVISION for s in structure.sections)
        has_toc = any(s.section_type == SectionType.TOC for s in structure.sections)

        lines.append(f"cover_enabled: {str(has_cover).lower()}")
        lines.append(f"signature_enabled: {str(has_signature).lower()}")
        lines.append(f"revision_enabled: {str(has_revision).lower()}")
        lines.append(f"toc_enabled: {str(has_toc).lower()}")

        # 扫描封面占位符
        cover_fields = self._scan_cover_placeholders(doc)
        if cover_fields:
            lines.append("cover_fields:")
            for k, v in cover_fields.items():
                lines.append(f"  {k}: \"{v}\"")
        return lines

    def _scan_cover_placeholders(self, doc) -> dict:
        """扫描封面区域的占位符 {{xxx}}"""
        fields = {}
        pattern = re.compile(r'\{\{(\w+)\}\}')
        for para in list(doc.paragraphs)[:30]:
            for m in pattern.finditer(para.text):
                key = m.group(1)
                if key not in fields:
                    fields[key] = ""
        return fields

    # ==================== Section 转换 ====================

    def _convert_section(self, section, doc, md_lines, images_dir, attachments_dir, docx_path):
        sect_type = section.section_type
        md_lines.append(f"<!-- section: {sect_type.value} -->")
        md_lines.append("")

        # 重新编号：进入新 section 时重置计数器
        self._section_fig_counter = 0
        self._section_tbl_counter = 0
        self._section_eq_counter = 0

        for elem in section.elements:
            para_index = elem.para_index
            if para_index < 0 or para_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_index]

            if elem.element_type == ElementType.HEADING:
                self._convert_heading(para, elem, md_lines, doc)

            elif elem.element_type == ElementType.IMAGE:
                self._convert_image(para, elem, md_lines, images_dir, attachments_dir, docx_path)

            elif elem.element_type == ElementType.EQUATION:
                self._convert_equation(para, elem, md_lines)

            elif elem.element_type == ElementType.TABLE:
                # 表格通过 doc.tables 按顺序处理
                pass

            else:
                # PARAGRAPH 类型
                self._convert_text_paragraph(para, md_lines, doc)

        md_lines.append("")

    # ==================== 书签/域/脚注 扫描 ====================

    def _scan_bookmarks(self, doc):
        """扫描文档中所有书签"""
        body = doc.element.body
        current_bookmark = None
        for elem in body.iter():
            if elem.tag == WML_NS + 'bookmarkStart':
                bm_id = elem.get(WML_NS + 'id', '')
                bm_name = elem.get(WML_NS + 'name', '')
                if bm_name:
                    self._bookmarks[bm_id] = {'name': bm_name, 'text': ''}
                    current_bookmark = bm_id
            elif elem.tag == WML_NS + 'bookmarkEnd':
                current_bookmark = None
            elif current_bookmark and elem.tag == WML_NS + 't':
                if current_bookmark in self._bookmarks:
                    self._bookmarks[current_bookmark]['text'] += (elem.text or '')

    def _scan_ref_fields(self, doc):
        """扫描所有 REF 域"""
        body = doc.element.body
        for instr in body.iter(WML_NS + 'instrText'):
            if instr.text and 'REF' in instr.text.upper():
                match = re.search(r'REF\s+(\S+)', instr.text, re.IGNORECASE)
                if match:
                    ref_name = match.group(1)
                    # 找到关联的显示文本
                    display_text = self._get_ref_display_text(instr)
                    self._ref_fields.append({
                        'bookmark': ref_name,
                        'display': display_text,
                    })

    def _get_ref_display_text(self, instr_elem) -> str:
        """获取 REF 域的显示文本"""
        parent = instr_elem.getparent()
        if parent is None:
            return ""
        grandparent = parent.getparent()
        if grandparent is None:
            return ""
        found_sep = False
        texts = []
        for child in grandparent.iter():
            if child.tag == WML_NS + 'fldChar':
                fld_type = child.get(WML_NS + 'fldCharType', '')
                if fld_type == 'separate':
                    found_sep = True
                elif fld_type == 'end':
                    break
            elif found_sep and child.tag == WML_NS + 't':
                texts.append(child.text or '')
        return ''.join(texts).strip()

    def _scan_footnotes(self, doc):
        """扫描脚注"""
        footnotes_part = None
        try:
            for rel in doc.part.rels.values():
                if 'footnote' in rel.reltype.lower():
                    footnotes_part = rel.target_part
                    break
        except Exception:
            pass

        if footnotes_part is None:
            return

        try:
            footnotes_xml = footnotes_part.blob.decode('utf-8', errors='replace')
            from lxml import etree
            root = etree.fromstring(footnotes_xml.encode('utf-8'))
            for fn_elem in root.iter(WML_NS + 'footnote'):
                fn_id = fn_elem.get(WML_NS + 'id', '')
                texts = []
                for t in fn_elem.iter(WML_NS + 't'):
                    texts.append(t.text or '')
                if fn_id and fn_id not in ('0', '-1'):
                    self._footnotes[fn_id] = ''.join(texts)
        except Exception as e:
            logger.debug(f"脚注扫描失败: {e}")

    def _append_footnote_definitions(self, md_lines):
        if not self._footnotes:
            return
        md_lines.append("")
        for fn_id in sorted(self._footnotes.keys(), key=int):
            fn_text = self._footnotes[fn_id]
            md_lines.append(f"[^{fn_id}]: {fn_text}")

    # ==================== 标题 ====================

    def _convert_heading(self, para, elem, md_lines, doc):
        level = elem.heading_level
        text = self._convert_inline_runs(para, doc)
        # 生成书签锚点
        bookmark_id = self._find_heading_bookmark(para)
        if bookmark_id:
            md_lines.append(f"{'#' * level} {text} {{#{bookmark_id}}}")
        else:
            md_lines.append(f"{'#' * level} {text}")
        md_lines.append("")

    def _find_heading_bookmark(self, para) -> str:
        """查找段落中的书签，用于生成锚点"""
        for elem in para._element.iter():
            if elem.tag == WML_NS + 'bookmarkStart':
                name = elem.get(WML_NS + 'name', '')
                if name and name != '_GoBack':
                    return f"sec:{name}"
        return ""

    # ==================== 图片 ====================

    def _convert_image(self, para, elem, md_lines, images_dir, attachments_dir, docx_path):
        self._section_fig_counter = self._section_fig_counter + 1
        fig_num = self._section_fig_counter

        # 提取图片
        img_filename = self._extract_image_from_para(para, docx_path, images_dir)
        alt_text = self._get_image_alt_text(para) or f"图片{fig_num}"

        # Visio OLE 处理
        visio_path = None
        if elem.text_preview and 'visio_ole' in elem.text_preview:
            visio_path = self._extract_visio_ole(para, docx_path, attachments_dir)

        if img_filename:
            rel_path = f"images/{img_filename}"
            md_lines.append(f"![{alt_text}]({rel_path})")
        else:
            md_lines.append(f"[{alt_text}]")

        # 题注文本
        caption_text = self._find_caption_for_para(para, "图") or alt_text
        # 书签锚点
        bookmark_id = self._find_para_bookmark(para)
        if bookmark_id:
            md_lines.append(f"*{caption_text}* {{#{bookmark_id}}}")
        else:
            md_lines.append(f"*{caption_text}*")

        if visio_path:
            md_lines.append(f"<!-- visio: attachments/{visio_path} -->")

        md_lines.append("")

    def _extract_image_from_para(self, para, docx_path, images_dir) -> str:
        """从段落中提取嵌入图片"""
        self.image_counter += 1
        for blip in para._element.iter(DML_NS + 'blip'):
            embed = blip.get(R_NS + 'embed', '')
            if embed:
                try:
                    rel = para.part.rels[embed]
                    image_part = rel.target_part
                    ext = Path(image_part.partname).suffix or '.png'
                    img_name = f"image_{self.image_counter}{ext}"
                    img_path = images_dir / img_name
                    with open(img_path, 'wb') as f:
                        f.write(image_part.blob)
                    return img_name
                except Exception as e:
                    logger.debug(f"图片提取失败: {e}")
        return ""

    def _get_image_alt_text(self, para) -> str:
        """获取图片的 alt 文本"""
        for elem in para._element.iter():
            if 'altText' in elem.tag or 'AltText' in elem.tag:
                return elem.get('val', '')
            if elem.tag.endswith('}docPr'):
                return elem.get('descr', '') or elem.get('name', '')
        return ""

    def _find_caption_for_para(self, para, prefix="图") -> str:
        """查找紧接图片后的题注段落"""
        body = para._element.getparent()
        if body is None:
            return ""
        paras = list(body.iterchildren(WML_NS + 'p'))
        try:
            idx = paras.index(para._element)
            if idx + 1 < len(paras):
                next_p = paras[idx + 1]
                text = ''.join(t.text or '' for t in next_p.iter(WML_NS + 't'))
                if text.strip():
                    return self._clean_caption_text(text.strip())
        except (ValueError, IndexError):
            pass
        return ""

    def _clean_caption_text(self, text: str) -> str:
        # 移除可能的前导编号（图序）保留描述
        text = re.sub(r'^(图|Figure|Fig\.?)\s*\d+[：:.\s]*', '', text).strip()
        text = re.sub(r'^(表|Table|Tbl\.?)\s*\d+[：:.\s]*', '', text).strip()
        return text

    # ==================== Visio OLE 提取 ====================

    def _extract_visio_ole(self, para, docx_path, attachments_dir) -> str:
        """提取 Visio OLE 二进制文件，返回文件名"""
        for obj in para._element.iter(WML_NS + 'object'):
            ole_obj = obj.find(OLE_NS + 'OLEObject')
            if ole_obj is not None:
                progid = ole_obj.get('ProgID', '')
                if 'Visio' in progid:
                    # 读取 OLE 对象原始数据
                    ole_id = ole_obj.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', '')
                    if ole_id:
                        try:
                            # OLE 数据存储在 docx zip 的 embeddings/ 目录
                            with zipfile.ZipFile(docx_path, 'r') as zf:
                                # 查找 oleObject 条目
                                for name in zf.namelist():
                                    if 'embeddings' in name.lower() and 'oleObject' in name.lower():
                                        # 匹配编号
                                        if ole_id.replace('rId', '') in name:
                                            ole_name = f"visio_{Path(name).stem}.vsd"
                                            ole_path = attachments_dir / ole_name
                                            with zf.open(name) as src:
                                                with open(ole_path, 'wb') as dst:
                                                    dst.write(src.read())
                                            return ole_name
                        except Exception as e:
                            logger.debug(f"Visio OLE 提取失败: {e}")
        return ""

    # ==================== 公式 ====================

    def _convert_equation(self, para, elem, md_lines):
        self._section_eq_counter = self._section_eq_counter + 1
        # 尝试提取 LaTeX
        latex = self._omml_to_latex(para)
        bookmark_id = self._find_para_bookmark(para)

        if latex:
            md_lines.append("$$")
            md_lines.append(latex)
            if bookmark_id:
                md_lines.append(f"$$ {{{bookmark_id}}}")
            else:
                md_lines.append("$$")
        else:
            # 降级：保留原文
            text = para.text.strip()
            if bookmark_id:
                md_lines.append(f"> [公式] {text} {{{bookmark_id}}}")
            else:
                md_lines.append(f"> [公式] {text}")

        md_lines.append("")

    def _omml_to_latex(self, para) -> str:
        """将 OMML 公式转换为 LaTeX（简化版）"""
        # 提取 OMML 文本内容
        parts = []
        for omath in para._element.iter(WML_NS + 'oMath'):
            for mt in omath.iter(WML_NS + 'm:t'):
                if mt.text:
                    parts.append(mt.text.strip())
        if parts:
            return ' '.join(parts)

        # 检查 MathType 图片中的 alt 文本
        for elem in para._element.iter(MC_NS + 'AlternateContent'):
            for t in elem.iter(WML_NS + 't'):
                if t.text and t.text.strip():
                    return t.text.strip()

        return ""

    # ==================== 文本段落 ====================

    def _convert_text_paragraph(self, para, md_lines, doc):
        text = para.text.strip()
        if not text:
            md_lines.append("")
            return

        # 检测脚注引用
        footnote_refs = self._find_footnote_refs_in_para(para, doc)

        # 转换内联格式
        md_text = self._convert_inline_runs(para, doc)

        # 添加脚注引用标记
        for fn_id in footnote_refs:
            md_text += f"[^{fn_id}]"

        # 检测列表
        if self._is_list_item(para):
            prefix = self._get_list_prefix(para)
            md_text = prefix + md_text

        md_lines.append(md_text)

    def _convert_inline_runs(self, para, doc) -> str:
        """转换段落中的行内格式（粗体、斜体、超链接、REF域）"""
        result = []
        for run in para.runs:
            text = run.text or ''
            if not text:
                continue

            # 检查 REF 域或 HYPERLINK
            ref_md = self._process_field_run(run, text)
            if ref_md:
                result.append(ref_md)
                continue

            if run.bold:
                text = f"**{text}**"
            if run.font.italic:
                text = f"*{text}*"

            result.append(text)

        return ''.join(result)

    def _process_field_run(self, run, text: str) -> str:
        """处理包含域的 run"""
        xml = run._element.xml
        text_clean = text.strip()

        # 检查是否是 REF 域的一部分
        for ref in self._ref_fields:
            if ref['display'] and ref['display'] in text_clean:
                bk_name = ref['bookmark']
                # 判断引用类型
                if bk_name.startswith('fig') or bk_name.startswith('figure'):
                    return text.replace(ref['display'], f"[{ref['display']}](#fig:{bk_name})")
                elif bk_name.startswith('tbl') or bk_name.startswith('table'):
                    return text.replace(ref['display'], f"[{ref['display']}](#tbl:{bk_name})")
                elif bk_name.startswith('eq') or bk_name.startswith('equation'):
                    return text.replace(ref['display'], f"[{ref['display']}](#eq:{bk_name})")
                else:
                    return text.replace(ref['display'], f"[{ref['display']}](#sec:{bk_name})")

        # 检查 HYPERLINK
        if 'HYPERLINK' in xml:
            match = re.search(r'HYPERLINK\s+"([^"]+)"', xml)
            if match:
                url = match.group(1)
                return f"[{text_clean}]({url})"

        # 检查 HYPERLINK 关系
        hyperlink = run._element.find(WML_NS + 'instrText')
        if hyperlink is not None and hyperlink.text and 'HYPERLINK' in hyperlink.text:
            match = re.search(r'HYPERLINK\s+"([^"]+)"', hyperlink.text)
            if match:
                url = match.group(1)
                return f"[{text_clean}]({url})"

        return ""

    def _find_footnote_refs_in_para(self, para, doc) -> list:
        """查找段落中的脚注引用"""
        refs = []
        for child in para._element.iter():
            if child.tag == WML_NS + 'footnoteReference':
                fn_id = child.get(WML_NS + 'id', '')
                if fn_id:
                    refs.append(fn_id)
        return refs

    # ==================== 表格 ====================

    def _convert_all_tables(self, doc, md_lines):
        """转换文档中的表格"""
        # 通过段落中嵌入的表格来处理
        for table in doc.tables:
            # 查找表格前可能的题注段落
            caption = self._find_caption_before_table(table, doc)
            md_lines.append("")
            if caption:
                bm_id = self._find_table_bookmark(table)
                if bm_id:
                    md_lines.append(f"*{caption}* {{#{bm_id}}}")
                else:
                    md_lines.append(f"*{caption}*")
                md_lines.append("")
            self._convert_table_to_md(table, md_lines)
            md_lines.append("")

    def _find_caption_before_table(self, table, doc) -> str:
        """查找表格前的题注段落"""
        tbl_elem = table._element
        body = tbl_elem.getparent()
        if body is None:
            return ""
        children = list(body)
        try:
            idx = children.index(tbl_elem)
            if idx > 0:
                prev = children[idx - 1]
                if prev.tag == WML_NS + 'p':
                    text = ''.join(t.text or '' for t in prev.iter(WML_NS + 't'))
                    if text.strip() and ('表' in text or 'Table' in text):
                        return text.strip()
        except ValueError:
            pass
        return ""

    def _find_table_bookmark(self, table) -> str:
        """查找表格元素中的书签"""
        for elem in table._element.iter():
            if elem.tag == WML_NS + 'bookmarkStart':
                name = elem.get(WML_NS + 'name', '')
                if name and name != '_GoBack':
                    return f"tbl:{name}"
        return ""

    def _convert_table_to_md(self, table, md_lines):
        rows = []
        max_cols = 0
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.replace('\n', ' ').strip()
                cells.append(cell_text)
            max_cols = max(max_cols, len(cells))
            rows.append(cells)

        if not rows:
            return

        # 确保所有行有相同列数
        for row in rows:
            while len(row) < max_cols:
                row.append("")

        # 表头
        md_lines.append("| " + " | ".join(rows[0]) + " |")
        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        # 数据行
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

    # ==================== 列表检测 ====================

    def _is_list_item(self, para) -> bool:
        pPr = para._element.find(WML_NS + 'pPr')
        if pPr is not None:
            numPr = pPr.find(WML_NS + 'numPr')
            if numPr is not None:
                return True
        # 也检查文本模式
        text = para.text.strip()
        if re.match(r'^[A-Z]\.\s', text) or re.match(r'^[a-z]\.\s', text):
            return True
        return False

    def _get_list_prefix(self, para) -> str:
        """获取列表前缀"""
        text = para.text.strip()
        # 检查字母编号
        m = re.match(r'^([A-Z])\.\s', text)
        if m:
            return f"{m.group(1)}. "
        m = re.match(r'^([a-z])\.\s', text)
        if m:
            return f"{m.group(1)}. "
        # 检查数字编号
        m = re.match(r'^(\d+)\.\s', text)
        if m:
            return f"{m.group(1)}. "
        return "- "

    # ==================== 书签辅助 ====================

    def _find_para_bookmark(self, para) -> str:
        """查找段落中的书签 ID"""
        for elem in para._element.iter():
            if elem.tag == WML_NS + 'bookmarkStart':
                name = elem.get(WML_NS + 'name', '')
                if name and name != '_GoBack':
                    prefix = "fig"
                    if 'tbl' in name.lower() or 'table' in name.lower():
                        prefix = "tbl"
                    elif 'eq' in name.lower() or 'equation' in name.lower():
                        prefix = "eq"
                    elif 'sec' in name.lower():
                        prefix = "sec"
                    return f"{prefix}:{name}"
        return ""

    # ==================== 便捷接口 ====================

    def convert_file(self, input_path: str, output_dir: str = None) -> str:
        """便捷接口：转换单个文件"""
        input_path = Path(input_path)
        if output_dir is None:
            output_dir = input_path.parent / "md_output"
        return self.convert(str(input_path), str(output_dir))


def convert_word_to_markdown(input_path: str, output_dir: str = None) -> str:
    """便捷函数"""
    converter = Word2MDConverter()
    return converter.convert_file(input_path, output_dir)
