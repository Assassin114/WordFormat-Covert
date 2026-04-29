"""
封面占位符替换器
支持：上下文感知替换、多个占位符匹配、签署页字段填充
"""

import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ..utils.logger import get_logger

logger = get_logger()

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class CoverReplacer:
    """
    封面占位符替换器
    
    支持两种替换模式：
    1. 精确替换：整个run的文本完全匹配占位符时替换
    2. 上下文替换：占位符嵌入在句子中间时替换
    """
    
    def __init__(self, fields: dict = None):
        """
        Args:
            fields: 字段字典，key为字段名，value为替换值
                   例如：{"contract_no": "ABC123", "classification": "秘密"}
        """
        self.fields = fields or {}
    
    def replace(self, doc: Document) -> int:
        """
        替换文档中的所有封面占位符
        
        Returns:
            替换的占位符数量
        """
        count = 0
        
        for para in doc.paragraphs:
            count += self._replace_paragraph(para)
        
        return count
    
    def _replace_paragraph(self, para) -> int:
        """替换段落中的占位符"""
        count = 0
        
        for run in para.runs:
            text = run.text
            if not text or not text.strip():
                continue
            
            replaced = self._replace_run_text(run, text)
            if replaced:
                count += replaced
        
        return count
    
    def _replace_run_text(self, run, text: str) -> int:
        """
        替换run中的文本
        返回替换的占位符数量
        """
        count = 0
        new_text = text
        
        for key, value in self.fields.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
                count += 1
                logger.debug(f"替换占位符: {placeholder} -> {value}")
        
        if count > 0:
            run.text = new_text
        
        return count
    
    def replace_table_fields(self, doc: Document) -> int:
        """
        替换表格中的占位符（用于签署页等）
        
        表格中的占位符通常是键值对形式：
        | 拟制：        | 时间：         |
        | {{drafter}}   | {{drafter_time}}|
        """
        count = 0
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                replaced = self._replace_run_text(run, run.text)
                                count += replaced
        
        return count


class CoverTemplate:
    """
    封面模板管理
    用于加载和管理封面模板的字段定义
    """
    
    # 标准封面字段定义
    STANDARD_FIELDS = {
        "contract_no": "合同编号",
        "classification": "密级",
        "year": "年份",
        "organization": "单位",
        "date": "日期",
        "drafter": "拟制",
        "drafter_time": "拟制时间",
        "reviewer": "校对",
        "reviewer_time": "校对时间",
        "approver": "标审",
        "approver_time": "标审时间",
        "checker": "审核",
        "checker_time": "审核时间",
        "final_approver": "批准",
        "final_approver_time": "批准时间",
    }
    
    @classmethod
    def create_fields_from_template(cls, template_path: str, value_map: dict = None) -> dict:
        """
        从封面模板创建字段字典
        
        Args:
            template_path: 模板文件路径（docx）
            value_map: 字段值映射，如果为None则使用占位符本身作为值
        
        Returns:
            字段字典
        """
        doc = Document(template_path)
        fields = {}
        
        # 扫描文档中的所有 {{xxx}} 占位符
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text
                if text:
                    placeholders = re.findall(r'\{\{(\w+)\}\}', text)
                    for key in placeholders:
                        if key not in fields:
                            if value_map and key in value_map:
                                fields[key] = value_map[key]
                            else:
                                fields[key] = cls.STANDARD_FIELDS.get(key, '{{' + key + '}}')

        # 也扫描表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text
                            if text:
                                placeholders = re.findall(r'\{\{(\w+)\}\}', text)
                                for key in placeholders:
                                    if key not in fields:
                                        if value_map and key in value_map:
                                            fields[key] = value_map[key]
                                        else:
                                            fields[key] = cls.STANDARD_FIELDS.get(key, '{{' + key + '}}')
        
        return fields
    
    @classmethod
    def get_field_list(cls) -> list:
        """获取所有标准字段的列表"""
        return list(cls.STANDARD_FIELDS.keys())


def replace_cover_placeholder(doc: Document, fields: dict) -> int:
    """
    便捷函数：替换封面占位符
    
    Args:
        doc: Document对象
        fields: 字段字典
    
    Returns:
        替换数量
    """
    replacer = CoverReplacer(fields)
    return replacer.replace(doc)
