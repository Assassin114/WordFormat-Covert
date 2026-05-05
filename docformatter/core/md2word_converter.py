"""
MD→Word 转换器 (Phase B-2)
将 Markdown 转换为格式化 Word 文档
接受标准 Markdown (L1) 或增强格式 (L2/L3)，所有特殊语法可选
"""

import re
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from ..utils.font_formatter import FontFormatter
from ..models.template_config import TemplateConfig, create_default_template
from ..utils.logger import get_logger

logger = get_logger()

WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class MD2WordConverter:
    """Markdown→Word 转换器"""

    def __init__(self, template: TemplateConfig = None):
        self.template = template or create_default_template()
        self._bookmark_registry: dict = {}
        self._ref_registry: list = []
        self._figure_count = 0
        self._table_count = 0
        self._equation_count = 0
        self._heading_count = 0
        self._footnote_counter = 0
        self._bookmark_id_counter = 0
        self._images_dir: Path = None

    def convert(self, md_path: str, output_path: str) -> str:
        """主入口：将 MD 文件转换为 Word 文档"""
        md_path = Path(md_path)
        md_content = md_path.read_text(encoding="utf-8")
        self._images_dir = md_path.parent / "images"
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        self._reset_state()

        # 解析 YAML frontmatter
        frontmatter, body_content = self._parse_frontmatter(md_content)

        # 应用 frontmatter 到模板
        if frontmatter:
            self._merge_frontmatter_to_template(frontmatter)

        # 解析 MD tokens
        tokens = self._parse_markdown(body_content)

        # 第一遍：预扫描，收集书签和编号信息
        self._prescan_tokens(tokens, md_path)

        # 创建 Word 文档
        doc = Document()

        # 应用页面设置
        self._apply_page_config(doc.sections[0])

        # 创建特殊页面
        if self.template.cover.enabled:
            self._create_cover_page(doc, frontmatter)
        if self.template.signature.enabled:
            self._create_signature_page(doc)
        if self.template.revision.enabled:
            self._create_revision_page(doc)
        if self.template.toc.enabled:
            self._create_toc_page(doc)

        # 处理正文 tokens
        self._process_tokens(tokens, doc, md_path)

        # 应用页眉页脚
        from .header_footer import HeaderFooterManager
        hf = HeaderFooterManager(self.template)
        hf.apply(doc)

        # 保存
        doc.save(str(output_path))
        logger.info(f"MD→Word 完成: {output_path}")
        return str(output_path)

    def _reset_state(self):
        self._bookmark_registry = {}
        self._ref_registry = []
        self._figure_count = 0
        self._table_count = 0
        self._equation_count = 0
        self._heading_count = 0
        self._footnote_counter = 0
        self._bookmark_id_counter = 0

    # ==================== Frontmatter 解析 ====================

    def _parse_frontmatter(self, content: str) -> tuple:
        """解析 YAML frontmatter，返回 (dict, body_text)"""
        frontmatter = {}
        body = content
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                try:
                    import yaml
                    frontmatter = yaml.safe_load(parts[1]) or {}
                except Exception:
                    frontmatter = self._parse_simple_frontmatter(parts[1])
                body = parts[2]
        return frontmatter, body.strip()

    def _parse_simple_frontmatter(self, text: str) -> dict:
        """简单解析 frontmatter（不依赖 yaml）"""
        result = {}
        current_key = None
        for line in text.strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            m = re.match(r"^(\w+):\s*(.*)", line)
            if m:
                key = m.group(1)
                val = m.group(2).strip().strip('"').strip("'")
                result[key] = val
                current_key = key
            elif current_key and line.startswith("- "):
                if current_key not in result:
                    result[current_key] = []
                if isinstance(result[current_key], list):
                    result[current_key].append(line[2:].strip().strip('"').strip("'"))
        return result

    def _merge_frontmatter_to_template(self, fm: dict):
        """将 frontmatter 设置合并到模板配置"""
        bool_keys = ["cover_enabled", "signature_enabled", "revision_enabled", "toc_enabled"]
        for key in bool_keys:
            if key in fm:
                val = fm[key]
                if isinstance(val, str):
                    val = val.lower() in ("true", "yes", "1")
                if key == "cover_enabled":
                    self.template.cover.enabled = val
                elif key == "signature_enabled":
                    self.template.signature.enabled = val
                elif key == "revision_enabled":
                    self.template.revision.enabled = val
                elif key == "toc_enabled":
                    self.template.toc.enabled = val

        if "cover_fields" in fm and isinstance(fm["cover_fields"], dict):
            for k, v in fm["cover_fields"].items():
                self.template.cover.fields[k] = str(v) if v else ""

    # ==================== MD 解析 ====================

    def _parse_markdown(self, content: str):
        """使用 markdown-it-py 解析 Markdown，预处理数学公式"""
        # 预存数学块 $$...$$ → fenced code block（在 token 处理中识别）
        self._math_blocks = {}
        math_counter = [0]

        def replace_display_math(m):
            math_counter[0] += 1
            key = f"__MATH_{math_counter[0]}__"
            self._math_blocks[key] = m.group(1).strip()
            return f"\n```math\n{key}\n```\n"

        content = re.sub(r"\$\$\s*(.+?)\s*\$\$", replace_display_math, content, flags=re.DOTALL)

        from markdown_it import MarkdownIt
        md = MarkdownIt("commonmark").enable(["table", "strikethrough"])
        return md.parse(content)

    # ==================== 预扫描 ====================

    def _prescan_tokens(self, tokens, md_path):
        """预扫描：收集书签信息"""
        i = 0
        while i < len(tokens):
            token = tokens[i]
            ttype = token.type

            if ttype == "heading_open":
                self._heading_count += 1
                next_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                text = next_tok.content if next_tok else f"heading_{self._heading_count}"
                bm = self._extract_bookmark_from_text(text)
                if bm:
                    self._bookmark_registry[bm] = {"text": self._strip_bm_syntax(text), "type": "heading"}
                else:
                    self._bookmark_registry[f"sec:heading_{self._heading_count}"] = {"text": text, "type": "heading"}

            elif ttype == "paragraph_open":
                # 检查段落内的 inline children（图片、题注等）
                inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                if inline_tok and inline_tok.type == "inline" and inline_tok.children:
                    for child in inline_tok.children:
                        if child.type == "image":
                            self._figure_count += 1
                            alt_text = child.content or f"图{self._figure_count}"
                            # 查找图片后的题注段落
                            caption_text = f"图 {self._figure_count}"
                            if i + 4 < len(tokens) and tokens[i + 3].type == "paragraph_open":
                                cap_inline = tokens[i + 4] if i + 4 < len(tokens) else None
                                if cap_inline and cap_inline.type == "inline" and cap_inline.content.strip().startswith("*"):
                                    caption_text = cap_inline.content.strip("*")
                            bm = self._extract_bookmark_from_text(caption_text)
                            if bm:
                                self._bookmark_registry[bm] = {"text": self._strip_bm_syntax(caption_text), "type": "fig"}
                            else:
                                self._bookmark_registry[f"fig:{self._figure_count}"] = {"text": caption_text, "type": "fig"}

            elif ttype == "table_open":
                self._table_count += 1
                caption_text = f"表 {self._table_count}"
                # 向前查找题注（*表注*）
                if i >= 2 and tokens[i - 2].type == "inline":
                    prev_text = tokens[i - 2].content.strip()
                    if prev_text.startswith("*"):
                        caption_text = prev_text.strip("*")
                bm = self._extract_bookmark_from_text(caption_text)
                if bm:
                    self._bookmark_registry[bm] = {"text": self._strip_bm_syntax(caption_text), "type": "tbl"}
                else:
                    self._bookmark_registry[f"tbl:{self._table_count}"] = {"text": caption_text, "type": "tbl"}

            elif ttype == "fence" and token.info.strip() == "math":
                # 数学块（已通过预处理转为 ```math fence）
                self._equation_count += 1
                self._bookmark_registry[f"eq:{self._equation_count}"] = {"text": f"公式 {self._equation_count}", "type": "eq"}

            i += 1

    def _extract_bookmark_from_text(self, text: str) -> str:
        """从文本中提取 {#xxx} 书签"""
        m = re.search(r"\{#([\w:.-]+)\}\s*$", text)
        if m:
            return m.group(1)
        return ""

    def _strip_bm_syntax(self, text: str) -> str:
        """移除文本中的 {#xxx} 标记"""
        return re.sub(r"\s*\{#[\w:.-]+\}\s*$", "", text).strip()

    # ==================== Token 处理主循环 ====================

    def _process_tokens(self, tokens, doc, md_path):
        i = 0
        section_type = "body"

        while i < len(tokens):
            token = tokens[i]
            ttype = token.type

            if ttype == "html_block":
                content = token.content.strip()
                section_match = re.search(r"<!--\s*section:\s*(\w+)\s*-->", content)
                if section_match:
                    section_type = section_match.group(1)
                    doc.add_section()
                    self._apply_page_config(doc.sections[-1])
                visio_match = re.search(r"<!--\s*visio:\s*(\S+)\s*-->", content)
                if visio_match:
                    visio_rel_path = visio_match.group(1)
                    visio_full = (md_path.parent / visio_rel_path).resolve()
                    if visio_full.exists():
                        self._embed_ole_object(doc, str(visio_full))
                i += 1
                continue

            elif ttype == "heading_open":
                i = self._process_heading_open(tokens, i, doc)

            elif ttype == "paragraph_open":
                i = self._process_paragraph_open(tokens, i, doc)

            elif ttype == "bullet_list_open" or ttype == "ordered_list_open":
                i = self._process_list(tokens, i, doc)

            elif ttype == "table_open":
                i = self._process_table(tokens, i, doc)

            elif ttype == "fence":
                i = self._process_code_block(token, doc)
                i += 1
                continue

            i += 1

    # ==================== 标题处理 ====================

    def _process_heading_open(self, tokens, start_idx, doc) -> int:
        token = tokens[start_idx]
        heading_tag = token.tag  # h1, h2, h3, h4, h5
        level = int(heading_tag[1]) if heading_tag.startswith("h") else 1

        inline_tok = tokens[start_idx + 1] if start_idx + 1 < len(tokens) else None
        text = inline_tok.content if inline_tok else ""

        bm = self._extract_bookmark_from_text(text)
        clean_text = self._strip_bm_syntax(text)

        para = doc.add_paragraph()
        run = para.add_run(clean_text)

        # 应用标题格式
        if 1 <= level <= len(self.template.headings):
            hc = self.template.headings[level - 1]
            FontFormatter.apply_font(run, hc.font)
            FontFormatter.apply_paragraph(para, hc.paragraph)

        # 创建书签
        if bm and bm in self._bookmark_registry:
            self._create_bookmark(para, bm)

        return start_idx + 3  # heading_open + inline + heading_close

    # ==================== 段落处理 ====================

    def _process_paragraph_open(self, tokens, start_idx, doc) -> int:
        inline_tok = tokens[start_idx + 1] if start_idx + 1 < len(tokens) else None
        if inline_tok is None or inline_tok.type != "inline":
            return start_idx + 3

        content = inline_tok.content.strip()

        # 检查是否仅包含图片
        if inline_tok.children and len(inline_tok.children) == 1 and inline_tok.children[0].type == "image":
            img_token = inline_tok.children[0]
            self._figure_count += 1
            alt = img_token.content or f"图{self._figure_count}"
            src = img_token.attrs.get("src", "") if img_token.attrs else ""

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_path = self._resolve_image_path(src, None)
            if img_path and img_path.exists():
                try:
                    from docx.opc.constants import RELATIONSHIP_TYPE as RT
                    with open(img_path, "rb") as f:
                        image_part = doc.part.get_or_add_image_part(f.read())
                    rId = doc.part.relate_to(image_part, RT.IMAGE)
                    run = para.add_run()
                    drawing_xml = self._make_inline_drawing(rId, width=Inches(5))
                    run._element.append(drawing_xml)
                except Exception as e:
                    logger.debug(f"图片插入失败: {e}")
                    para.add_run(f"[图片: {alt}]")
            else:
                para.add_run(f"[图片缺失: {alt}]")

            doc.add_paragraph()
            return start_idx + 3

        if not content:
            return start_idx + 3

        para = doc.add_paragraph()
        FontFormatter.apply_paragraph(para, self.template.body.paragraph)

        # 直接渲染简单段落
        if inline_tok.children:
            self._render_inline_tokens(para, inline_tok.children, doc)
        else:
            # 处理内联格式和引用
            self._render_paragraph_text_with_refs(para, content)

        return start_idx + 3  # p_open + inline + p_close

    def _render_paragraph_text_with_refs(self, para, text: str):
        """渲染段落文本，处理 [text](#id) 引用"""
        # 分割：普通文本 vs [text](#id) vs [text](url)
        pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
        last_end = 0
        for m in pattern.finditer(text):
            # 前面的普通文本
            before = text[last_end:m.start()]
            if before:
                run = para.add_run(before)
                FontFormatter.apply_font(run, self.template.body.font)

            link_text = m.group(1)
            link_target = m.group(2)

            if link_target.startswith("#"):
                # 内部引用 → REF 域
                ref_id = link_target[1:]  # 去掉 #
                if ref_id in self._bookmark_registry:
                    self._create_ref_field(para, link_text, ref_id)
                else:
                    run = para.add_run(link_text)
                    FontFormatter.apply_font(run, self.template.body.font)
            else:
                # 外部超链接 → HYPERLINK
                self._create_hyperlink(para, link_text, link_target)

            last_end = m.end()

        # 剩余文本
        remaining = text[last_end:]
        if remaining:
            # 处理脚注 [^id]
            remaining = self._render_text_with_footnotes(para, remaining)

            run = para.add_run(remaining)
            FontFormatter.apply_font(run, self.template.body.font)

    def _render_inline_tokens(self, para, children, doc):
        """从 markdown-it inline children 渲染段落，支持所有内联格式"""
        i = 0
        while i < len(children):
            child = children[i]
            ctype = child.type

            if ctype == "text":
                run = para.add_run(child.content)
                FontFormatter.apply_font(run, self.template.body.font)

            elif ctype == "strong_open":
                i += 1
                while i < len(children) and children[i].type != "strong_close":
                    if children[i].type == "text":
                        run = para.add_run(children[i].content)
                        run.bold = True
                        FontFormatter.apply_font(run, self.template.body.font)
                    i += 1

            elif ctype == "em_open":
                i += 1
                while i < len(children) and children[i].type != "em_close":
                    if children[i].type == "text":
                        run = para.add_run(children[i].content)
                        run.italic = True
                        FontFormatter.apply_font(run, self.template.body.font)
                    i += 1

            elif ctype == "code_inline":
                run = para.add_run(child.content)
                FontFormatter.apply_font(run, self.template.code_block.font)

            elif ctype == "link_open":
                href = child.attrs.get("href", "") if child.attrs else ""
                i += 1
                link_text_parts = []
                while i < len(children) and children[i].type != "link_close":
                    if children[i].type == "text":
                        link_text_parts.append(children[i].content)
                    i += 1
                link_text = "".join(link_text_parts)

                if href.startswith("#"):
                    ref_id = href[1:]
                    if ref_id in self._bookmark_registry:
                        self._create_ref_field(para, link_text, ref_id)
                    else:
                        run = para.add_run(link_text)
                        FontFormatter.apply_font(run, self.template.body.font)
                elif href:
                    self._create_hyperlink(para, link_text, href)
                else:
                    run = para.add_run(link_text)
                    FontFormatter.apply_font(run, self.template.body.font)

            elif ctype == "softbreak":
                para.add_run(" ")

            elif ctype == "hardbreak":
                para.add_run("\n")

            i += 1

    def _render_text_with_footnotes(self, para, text: str) -> str:
        """处理文本中的 [^id] 脚注引用，返回清除后的文本"""
        return re.sub(r"\[\^(\d+)\]", "", text)

    # ==================== 表格 ====================

    def _process_table(self, tokens, start_idx, doc) -> int:
        # 查找 caption
        caption_text = ""
        bm = ""
        if start_idx >= 2 and tokens[start_idx - 2].type == "inline":
            prev_content = tokens[start_idx - 2].content.strip()
            if prev_content.startswith("*"):
                caption_text = prev_content.strip("*")
                bm = self._extract_bookmark_from_text(caption_text)
                if bm:
                    caption_text = self._strip_bm_syntax(caption_text)

        if not bm:
            bm = f"tbl:{self._table_count}"
            if bm not in self._bookmark_registry:
                bm = ""

        # 收集表格数据
        rows_data = []
        i = start_idx + 1  # skip table_open
        thead_open = False
        tbody_open = False
        in_row = False
        current_row = []

        max_i = min(start_idx + 100, len(tokens))
        while i < max_i:
            tok = tokens[i]
            ttype = tok.type

            if ttype == "thead_open":
                thead_open = True
            elif ttype == "thead_close":
                thead_open = False
            elif ttype == "tbody_open":
                tbody_open = True
            elif ttype == "tbody_close":
                tbody_open = False
            elif ttype == "tr_open":
                in_row = True
                current_row = []
            elif ttype == "tr_close":
                in_row = False
                rows_data.append(current_row)
            elif ttype in ("th_open", "td_open"):
                inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                cell_text = inline_tok.content if inline_tok and inline_tok.type == "inline" else ""
                current_row.append(cell_text)
                i += 2  # skip td_open + inline
                continue
            elif ttype == "table_close":
                break

            i += 1

        if not rows_data:
            return i

        # 创建 Word 表格
        max_cols = max(len(r) for r in rows_data)
        for r in rows_data:
            while len(r) < max_cols:
                r.append("")

        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(rows_data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                if row_idx == 0:
                    FontFormatter.apply_font(run, self.template.table_font.regular_table.header_font)
                else:
                    FontFormatter.apply_font(run, self.template.table_font.regular_table.body_font)

        doc.add_paragraph()  # table 后空行

        return i

    # ==================== 代码块 ====================

    def _process_code_block(self, token, doc) -> int:
        info = token.info.strip() if token.info else ""

        # 数学块检测
        if info == "math":
            key = token.content.strip()
            latex_source = self._math_blocks.get(key, key)
            self._equation_count += 1

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            png_bytes = self._render_latex_to_png(latex_source)
            if png_bytes:
                try:
                    from docx.opc.constants import RELATIONSHIP_TYPE as RT
                    image_part = doc.part.get_or_add_image_part(png_bytes)
                    rId = doc.part.relate_to(image_part, RT.IMAGE)
                    run = para.add_run()
                    drawing_xml = self._make_inline_drawing(rId, width=Inches(4))
                    run._element.append(drawing_xml)
                except Exception as e:
                    logger.debug(f"公式图片插入失败: {e}")
                    run = para.add_run(f"[公式: {latex_source}]")
            else:
                run = para.add_run(f"[公式: {latex_source}]")
                FontFormatter.apply_font(run, self.template.body.font)

            para = doc.add_paragraph()  # 公式后空行
            return 0

        # 普通代码块
        code_text = token.content
        lines = code_text.split("\n")

        for line in lines:
            para = doc.add_paragraph()
            FontFormatter.apply_paragraph(para, self.template.code_block.paragraph)
            run = para.add_run(line or " ")
            FontFormatter.apply_font(run, self.template.code_block.font)

            # 设置底纹
            if self.template.code_block.bg_color:
                try:
                    from docx.oxml import parse_xml
                    color = self.template.code_block.bg_color.lstrip("#")
                    shading_xml = f'<w:shd {qn("w:fill")}="{color}" {qn("w:val")}="clear" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    pPr = para._element.get_or_add_pPr()
                    pPr.append(parse_xml(shading_xml))
                except Exception:
                    pass

        return 0  # 返回无意义，调用者用 continue

    def _render_latex_to_png(self, latex: str) -> bytes:
        """使用 matplotlib 渲染 LaTeX 为 PNG，返回 bytes"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            fig, ax = plt.subplots(figsize=(0.1, 0.1))
            text_obj = ax.text(0.5, 0.5, f"${latex}$", fontsize=12,
                               horizontalalignment="center", verticalalignment="center",
                               transform=ax.transAxes)
            ax.axis("off")

            # 计算边界
            fig.canvas.draw()
            bbox = text_obj.get_window_extent(renderer=fig.canvas.get_renderer())
            bbox = bbox.transformed(fig.dpi_scale_trans.inverted())
            plt.close(fig)

            # 重新绘制，使用正确的 figsize
            pad = 0.1
            fig, ax = plt.subplots(figsize=(bbox.width + pad, bbox.height + pad))
            ax.text(0.5, 0.5, f"${latex}$", fontsize=12,
                    horizontalalignment="center", verticalalignment="center",
                    transform=ax.transAxes)
            ax.axis("off")

            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", pad_inches=0.05)
            plt.close(fig)
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.debug(f"LaTeX 渲染失败: {e}")
            return b""

    def _resolve_image_path(self, src: str, md_path=None) -> Path:
        """解析图片路径"""
        if not src:
            return None
        candidates = []
        if md_path is not None:
            candidates.append(md_path.parent / src)
        if self._images_dir:
            candidates.append(self._images_dir / Path(src).name)
        for p in candidates:
            if p and p.exists():
                return p
        return None

    def _make_inline_drawing(self, rId, width=None, height=None) -> OxmlElement:
        """创建 inline drawing XML 元素"""
        drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        inline_xml = f"""<wp:inline xmlns:wp="{drawing_ns}" xmlns:a="{a_ns}" xmlns:r="{r_ns}" xmlns:pic="{pic_ns}" distT="0" distB="0" distL="0" distR="0">
    <wp:extent cx="{int(width * 914400)}" cy="{int((width or Inches(1)) * 600000)}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="1" name="picture"/>
    <a:graphic>
        <a:graphicData uri="{pic_ns}">
            <pic:pic>
                <pic:nvPicPr>
                    <pic:cNvPr id="0" name="picture"/>
                    <pic:cNvPicPr/>
                </pic:nvPicPr>
                <pic:blipFill>
                    <a:blip r:embed="{rId}"/>
                    <a:stretch>
                        <a:fillRect/>
                    </a:stretch>
                </pic:blipFill>
                <pic:spPr>
                    <a:xfrm>
                        <a:off x="0" y="0"/>
                        <a:ext cx="{int(width * 914400)}" cy="{int((width or Inches(1)) * 600000)}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"/>
                </pic:spPr>
            </pic:pic>
        </a:graphicData>
    </a:graphic>
</wp:inline>"""
        return etree.fromstring(inline_xml.encode("utf-8"))

    # ==================== 列表 ====================

    def _process_list(self, tokens, start_idx, doc) -> int:
        i = start_idx
        list_type = tokens[i].type  # bullet_list_open or ordered_list_open
        close_type = list_type.replace("_open", "_close")
        item_counter = 0

        i += 1
        while i < len(tokens) and tokens[i].type != close_type:
            tok = tokens[i]
            if tok.type == "list_item_open":
                i += 1
                item_counter += 1
                # 跳过 paragraph_open
                if i < len(tokens) and tokens[i].type == "paragraph_open":
                    i += 1
                # 读取 inline content
                if i < len(tokens) and tokens[i].type == "inline":
                    content = tokens[i].content
                    i += 1

                    para = doc.add_paragraph()
                    FontFormatter.apply_paragraph(para, self.template.body.paragraph)

                    prefix = f"{item_counter}. " if list_type == "ordered_list_open" else "- "
                    run = para.add_run(prefix + content)
                    FontFormatter.apply_font(run, self.template.body.font)

                # 跳过 paragraph_close, list_item_close
                if i < len(tokens) and tokens[i].type == "paragraph_close":
                    i += 1
                if i < len(tokens) and tokens[i].type == "list_item_close":
                    i += 1
            else:
                i += 1

        return i + 1  # 跳过 close 标签

    # ==================== 书签 ====================

    def _create_bookmark(self, para, bookmark_name: str):
        """在段落中创建 Word 书签，包裹整个段落内容"""
        self._bookmark_id_counter += 1
        bm_id = str(self._bookmark_id_counter)

        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(WML_NS + "id", bm_id)
        bookmark_start.set(WML_NS + "name", bookmark_name)

        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(WML_NS + "id", bm_id)

        if para.runs:
            first_run = para.runs[0]._element
            last_run = para.runs[-1]._element
            first_run.addprevious(bookmark_start)
            last_run.addnext(bookmark_end)
        else:
            run = para.add_run("")
            run._element.addprevious(bookmark_start)
            run._element.addnext(bookmark_end)

    # ==================== REF 域 ====================

    def _create_ref_field(self, para, display_text: str, bookmark_name: str):
        """创建 REF 域（可点击跳转的交叉引用）"""
        try:
            run_begin = para.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(WML_NS + "fldCharType", "begin")
            run_begin._element.append(fld_begin)

            run_instr = para.add_run()
            instr = OxmlElement("w:instrText")
            instr.set(WML_NS + "xml:space", "preserve")
            instr.text = f" REF {bookmark_name} \\h "
            run_instr._element.append(instr)

            run_sep = para.add_run()
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(WML_NS + "fldCharType", "separate")
            run_sep._element.append(fld_sep)

            run_text = para.add_run(display_text)
            FontFormatter.apply_font(run_text, self.template.body.font)

            run_end = para.add_run()
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(WML_NS + "fldCharType", "end")
            run_end._element.append(fld_end)
        except Exception as e:
            logger.debug(f"REF 域创建失败: {e}")
            run = para.add_run(display_text)
            FontFormatter.apply_font(run, self.template.body.font)

    # ==================== HYPERLINK ====================

    def _create_hyperlink(self, para, display_text: str, url: str):
        """创建外部超链接"""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            hyperlink_part = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
            rId = hyperlink_part.rId

            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(WML_NS + "history", "1")
            hyperlink.set(qn("r:id"), rId)

            run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(WML_NS + "val", "Hyperlink")
            rPr.append(rStyle)
            color = OxmlElement("w:color")
            color.set(WML_NS + "val", "0563C1")
            rPr.append(color)
            u = OxmlElement("w:u")
            u.set(WML_NS + "val", "single")
            rPr.append(u)
            run.append(rPr)

            t = OxmlElement("w:t")
            t.set(WML_NS + "xml:space", "preserve")
            t.text = display_text
            run.append(t)
            hyperlink.append(run)

            para._element.append(hyperlink)
        except Exception as e:
            logger.debug(f"HYPERLINK 创建失败: {e}")
            run = para.add_run(display_text)
            FontFormatter.apply_font(run, self.template.body.font)

    # ==================== 页面设置 ====================

    def _apply_page_config(self, section):
        """将模板的页面设置应用到指定节"""
        pc = self.template.page
        if pc.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(pc.height)
            section.page_height = Cm(pc.width)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(pc.width)
            section.page_height = Cm(pc.height)

        section.top_margin = Mm(pc.margin_top)
        section.bottom_margin = Mm(pc.margin_bottom)
        section.left_margin = Mm(pc.margin_left)
        section.right_margin = Mm(pc.margin_right)
        section.gutter = Mm(pc.gutter)

    # ==================== 特殊页面 ====================

    def _create_cover_page(self, doc, frontmatter: dict):
        """创建封面页"""
        from .cover_replacer import CoverReplacer, CoverTemplate

        # 如果有封面模板，从模板创建
        template_path = self.template.cover.template_path
        if template_path and Path(template_path).exists():
            try:
                cover_doc = Document(template_path)
                # 复制封面内容到主文档
                for para in cover_doc.paragraphs:
                    new_para = doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        if run.bold:
                            new_run.bold = True
                        if run.font.size:
                            new_run.font.size = run.font.size
                doc.add_page_break()
                return
            except Exception as e:
                logger.warning(f"封面模板加载失败: {e}")

        # 否则创建默认封面
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("文档标题")
        FontFormatter.apply_font(run, self.template.cover.heading_font)
        FontFormatter.apply_paragraph(para, self.template.cover.heading_paragraph)

        fields = self.template.cover.fields
        if fields:
            for key, val in fields.items():
                para = doc.add_paragraph()
                run = para.add_run(f"{key}: {val}")
                FontFormatter.apply_font(run, self.template.body.font)

        doc.add_page_break()

    def _create_signature_page(self, doc):
        """创建签署页"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.template.signature.title)
        FontFormatter.apply_font(run, self.template.signature.heading_font)
        FontFormatter.apply_paragraph(para, self.template.signature.heading_paragraph)

        # 创建简易签署表格
        sig_fields = ["拟制", "校对", "标审", "审核", "批准"]
        table = doc.add_table(rows=len(sig_fields), cols=3)
        table.style = "Table Grid"
        for i, field in enumerate(sig_fields):
            table.cell(i, 0).text = field
            table.cell(i, 1).text = "___________"
            table.cell(i, 2).text = "日期: ___________"

        doc.add_page_break()

    def _create_revision_page(self, doc):
        """创建修改记录页"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("修改记录")
        FontFormatter.apply_font(run, self.template.revision.heading_font)
        FontFormatter.apply_paragraph(para, self.template.revision.heading_paragraph)

        headers = ["序号", "版本", "日期", "修改内容", "作者"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            p = cell.paragraphs[0]
            for run in p.runs:
                FontFormatter.apply_font(run, self.template.revision.table_header_font)

        # 添加几个空行
        for _ in range(3):
            table.add_row()

        doc.add_page_break()

    def _create_toc_page(self, doc):
        """创建目录页"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("目录")
        FontFormatter.apply_font(run, self.template.toc.heading_font)
        FontFormatter.apply_paragraph(para, self.template.toc.heading_paragraph)

        # 插入 TOC 域
        para = doc.add_paragraph()
        run_begin = para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(WML_NS + "fldCharType", "begin")
        run_begin._element.append(fld_begin)

        run_instr = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(WML_NS + "xml:space", "preserve")
        instr.text = " TOC \\o \"1-5\" \\h \\z "
        run_instr._element.append(instr)

        run_sep = para.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(WML_NS + "fldCharType", "separate")
        run_sep._element.append(fld_sep)

        run_text = para.add_run("（请在 Word 中右键更新目录）")
        FontFormatter.apply_font(run_text, self.template.toc.entry_font)

        run_end = para.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(WML_NS + "fldCharType", "end")
        run_end._element.append(fld_end)

        doc.add_page_break()

    # ==================== OLE 嵌入 ====================

    def _embed_ole_object(self, doc, ole_path: str):
        """嵌入 OLE 对象（Visio 等）"""
        logger.info(f"OLE 嵌入: {ole_path}")
        # OLE 嵌入非常复杂，需要构造完整的 OLE object XML
        # 此处留作占位，实际实现需要操作 OOXML 底层
        para = doc.add_paragraph(f"[OLE 对象: {Path(ole_path).name}]")
