"""
核心格式化引擎
负责调度各子模块完成文档格式化
"""

from typing import List, Optional, Callable
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lxml import etree

from ..models.template_config import (
    TemplateConfig, FormatResult, BatchResult, 
    PrintMode, NumberingMode
)
from ..utils import (
    set_paragraph_properties,
    set_run_font,
    set_widow_control,
    set_keep_with_next,
    insert_page_break,
    is_landscape_section,
    get_paragraph_pPr,
)
from ..utils.logger import get_logger
from .table_handler import TableHandler
from .cover_replacer import CoverReplacer

logger = get_logger()

# WordprocessingML namespace
WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class DocumentFormatter:
    """
    核心格式化引擎，负责调度各子模块完成文档格式化
    """
    
    def __init__(self, template: TemplateConfig, cover_template_path: Optional[str] = None):
        self.template = template
        self.cover_template_path = cover_template_path
        self.style_mapper = StyleMapper(template)
        self.numbering_mgr = NumberingManager(template.caption)
        self.toc_gen = TOCGenerator()
        self.cross_ref_mgr = CrossReferenceManager()
        self.table_handler = TableHandler()  # 集成表格处理器
    
    def format(self, input_path: str, output_path: str) -> FormatResult:
        """
        格式化单个文档
        """
        result = FormatResult(
            success=False,
            input_path=input_path,
            output_path=output_path
        )
        
        try:
            logger.info(f"开始格式化: {input_path}")
            
            # Step 1: 解析文档
            doc = Document(input_path)
            
            # Step 2: 封面替换
            if self.cover_template_path and Path(self.cover_template_path).exists():
                self._apply_cover(doc)
            
            # Step 3: 样式映射
            style_mapping = self.style_mapper.analyze(doc)
            
            # Step 3.5: 扫描交叉引用
            self.cross_ref_mgr.scan_document(doc)
            
            # Step 4: 预扫描编号
            counts = self.numbering_mgr.process_document(doc)
            result.figures_processed = counts['figures']
            result.tables_processed = counts['tables']
            result.equations_processed = counts['equations']
            
            # Step 5: 应用格式到所有段落
            self._apply_formats(doc, style_mapping)
            
            # Step 6: 处理图片（添加图序、同页控制）
            result.figures_processed = self._process_images(doc)
            
            # Step 7: 处理表格（添加表序、同页控制） + 应用表格格式
            result.tables_processed = self._process_tables(doc)
            
            # Step 8: 处理公式（编号、居中）
            result.equations_processed = self._process_equations(doc)
            
            # Step 9: 生成目录
            self.toc_gen.generate_toc(doc, self.template.headings)
            
            # Step 10: 应用页眉页脚设置
            self._apply_header_footer(doc)
            
            # Step 11: 处理打印模式
            self._handle_print_mode(doc)
            
            # Step 12: 保存
            doc.save(output_path)
            result.success = True
            logger.info(f"格式化完成: {output_path}")
            
        except Exception as e:
            logger.error(f"格式化失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            result.errors.append(str(e))
        
        return result
    
    def batch_format(self, 
                     input_paths: List[str], 
                     output_dir: str,
                     progress_callback: Optional[Callable[[int, int, str], None]] = None) -> BatchResult:
        """
        批量格式化文档
        """
        results = []
        total = len(input_paths)
        
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        for i, input_path in enumerate(input_paths):
            input_name = Path(input_path).stem
            output_path = str(Path(output_dir) / f"{input_name}_formatted.docx")
            
            if progress_callback:
                progress_callback(i + 1, total, f"正在处理: {input_name}")
            
            result = self.format(input_path, output_path)
            results.append(result)
        
        success_count = sum(1 for r in results if r.success)
        failed_count = total - success_count
        
        if progress_callback:
            progress_callback(total, total, f"完成！成功: {success_count}, 失败: {failed_count}")
        
        return BatchResult(
            total=total,
            success_count=success_count,
            failed_count=failed_count,
            results=results
        )
    
    def _apply_cover(self, doc: Document):
        """替换封面占位符"""
        if not self.template.cover.fields:
            return
        
        replacer = CoverReplacer(self.template.cover.fields)
        count = replacer.replace(doc)
        logger.info(f"封面占位符替换完成，共替换 {count} 处")
        
        # 同时处理表格中的签署页字段
        count += replacer.replace_table_fields(doc)
        logger.info(f"签署页字段替换完成，共替换 {count} 处")
    
    def _apply_formats(self, doc, style_mapping: dict):
        """应用格式到所有段落"""
        for para in doc.paragraphs:
            heading_level = self.style_mapper.get_heading_level(para)
            
            if heading_level is not None and heading_level <= len(self.template.headings):
                heading_config = self.template.headings[heading_level - 1]
                self._apply_heading_format(para, heading_config)
            else:
                self._apply_body_format(para, self.template.body)
    
    def _apply_heading_format(self, para, heading_config):
        """应用标题格式"""
        set_paragraph_properties(
            para,
            line_spacing=heading_config.paragraph.line_spacing,
            space_before=heading_config.paragraph.space_before,
            space_after=heading_config.paragraph.space_after,
            alignment=heading_config.paragraph.alignment,
        )
        
        for run in para.runs:
            set_run_font(
                run,
                font_name=heading_config.font.name,
                font_size=heading_config.font.size,
                bold=heading_config.font.bold,
                italic=heading_config.font.italic,
                color=heading_config.font.color,
            )
        
        para.style = f'Heading {heading_config.level}'
    
    def _apply_body_format(self, para, body_config):
        """应用正文格式"""
        set_paragraph_properties(
            para,
            line_spacing=body_config.line_spacing,
            space_before=body_config.space_before,
            space_after=body_config.space_after,
            first_line_indent=body_config.first_line_indent,
            alignment=body_config.alignment,
        )
        
        for run in para.runs:
            set_run_font(
                run,
                font_name=body_config.font.name,
                font_size=body_config.font.size,
                bold=body_config.font.bold,
                italic=body_config.font.italic,
                color=body_config.font.color,
            )
    
    def _process_images(self, doc: Document) -> int:
        """
        处理所有图片：添加图序、同页控制
        """
        count = 0
        figure_num = 0
        
        for para in doc.paragraphs:
            if self._paragraph_has_image(para):
                figure_num += 1
                caption_text = self.numbering_mgr.format_figure_caption(figure_num)
                self._insert_caption_below(para, caption_text, is_figure=True)
                count += 1
        
        return count
    
    def _paragraph_has_image(self, para) -> bool:
        """检查段落是否包含图片"""
        drawings = para._element.findall('.//' + WML_NS + 'drawing')
        if drawings:
            return True
        
        inlines = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        anchors = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
        
        return len(inlines) > 0 or len(anchors) > 0
    
    def _insert_caption_below(self, para, caption_text: str, is_figure: bool = True):
        """在段落下方插入题注"""
        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        new_para.append(pPr)
        
        jc = OxmlElement('w:jc')
        jc.set(WML_NS + 'val', 'center')
        pPr.append(jc)
        
        widowControl = OxmlElement('w:widowControl')
        pPr.append(widowControl)
        
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)
        
        spacing = OxmlElement('w:spacing')
        spacing.set(WML_NS + 'before', '120')
        pPr.append(spacing)
        
        run = OxmlElement('w:r')
        new_para.append(run)
        
        rPr = OxmlElement('w:rPr')
        run.append(rPr)
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(WML_NS + 'ascii', '宋体')
        rFonts.set(WML_NS + 'hAnsi', '宋体')
        rFonts.set(WML_NS + 'eastAsia', '宋体')
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(WML_NS + 'val', '21')
        rPr.append(sz)
        
        szCs = OxmlElement('w:szCs')
        szCs.set(WML_NS + 'val', '21')
        rPr.append(szCs)
        
        t = OxmlElement('w:t')
        t.text = caption_text
        run.append(t)
        
        parent = para._element.getparent()
        idx = list(parent).index(para._element)
        parent.insert(idx + 1, new_para)
    
    def _process_tables(self, doc: Document) -> int:
        """
        处理所有表格：添加表序、同页控制 + 应用表格格式
        """
        count = 0
        table_num = 0
        
        for table in doc.tables:
            table_num += 1
            caption_text = self.numbering_mgr.format_table_caption(table_num)
            
            # 在表格上方插入题注
            self._insert_caption_above_table(table, caption_text)
            count += 1
            
            # 应用表格格式化（使用TableHandler）
            table_type = self.table_handler.detect_table_type(table)
            self.table_handler.format_table(table, table_type)
        
        return count
    
    def _insert_caption_above_table(self, table, caption_text: str):
        """在表格上方插入题注"""
        tbl_element = table._element
        parent = tbl_element.getparent()
        
        if parent is None:
            return
        
        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        new_para.append(pPr)
        
        jc = OxmlElement('w:jc')
        jc.set(WML_NS + 'val', 'center')
        pPr.append(jc)
        
        widowControl = OxmlElement('w:widowControl')
        pPr.append(widowControl)
        
        spacing = OxmlElement('w:spacing')
        spacing.set(WML_NS + 'after', '120')
        pPr.append(spacing)
        
        run = OxmlElement('w:r')
        new_para.append(run)
        
        rPr = OxmlElement('w:rPr')
        run.append(rPr)
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(WML_NS + 'ascii', '宋体')
        rFonts.set(WML_NS + 'hAnsi', '宋体')
        rFonts.set(WML_NS + 'eastAsia', '宋体')
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(WML_NS + 'val', '21')
        rPr.append(sz)
        
        szCs = OxmlElement('w:szCs')
        szCs.set(WML_NS + 'val', '21')
        rPr.append(szCs)
        
        t = OxmlElement('w:t')
        t.text = caption_text
        run.append(t)
        
        tbl_idx = list(parent).index(tbl_element)
        parent.insert(tbl_idx, new_para)
    
    def _process_equations(self, doc: Document) -> int:
        """处理所有公式：编号、居中"""
        count = 0
        equation_num = 0
        
        for para in doc.paragraphs:
            if self._is_equation_para(para):
                equation_num += 1
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_text = self.numbering_mgr.format_equation_caption(equation_num)
                self._add_equation_number(para, caption_text)
                count += 1
        
        return count
    
    def _is_equation_para(self, para) -> bool:
        """判断段落是否为公式"""
        omath_elements = para._element.findall('.//' + WML_NS + 'oMath')
        if omath_elements:
            return True
        
        for run in para.runs:
            try:
                rPr = run._element.find(WML_NS + 'rPr')
                if rPr is not None:
                    mathPr = rPr.find('.//' + WML_NS + 'oMath')
                    if mathPr is not None:
                        return True
            except Exception:
                pass
        
        return False
    
    def _add_equation_number(self, para, equation_num: str):
        """在段落右侧添加工公式编号"""
        pPr = get_paragraph_pPr(para)
        
        tabs = pPr.find(WML_NS + 'tabs')
        if tabs is None:
            tabs = OxmlElement('w:tabs')
            pPr.append(tabs)
        
        for tab in tabs.findall(WML_NS + 'tab'):
            tabs.remove(tab)
        
        tab = OxmlElement('w:tab')
        tab.set(WML_NS + 'val', 'right')
        tab.set(WML_NS + 'leader', 'none')
        tab.set(WML_NS + 'pos', '8298')
        tabs.append(tab)
        
        run = OxmlElement('w:r')
        para._element.append(run)
        
        tab_char = OxmlElement('w:tab')
        run.append(tab_char)
        
        run2 = OxmlElement('w:r')
        para._element.append(run2)
        
        rPr = OxmlElement('w:rPr')
        run2.append(rPr)
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(WML_NS + 'ascii', '宋体')
        rFonts.set(WML_NS + 'hAnsi', '宋体')
        rFonts.set(WML_NS + 'eastAsia', '宋体')
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(WML_NS + 'val', '21')
        rPr.append(sz)
        
        t = OxmlElement('w:t')
        t.text = equation_num
        run2.append(t)
    
    def _apply_header_footer(self, doc: Document):
        """应用页眉页脚设置"""
        for section in doc.sections:
            if self.template.header_footer.show_on_first:
                section.different_first_page_header_footer = True
            
            if self.template.header_footer.different_odd_even:
                section.even_odd_header_footer = True
    
    def _handle_print_mode(self, doc: Document):
        """处理打印模式 - 双面打印时确保文档从偶数页开始"""
        if self.template.print_mode == PrintMode.DUPLEX:
            body = doc.element.body
            
            last_elem = list(body)[-1] if list(body) else None
            if last_elem is None:
                return
            
            page_break_para = OxmlElement('w:p')
            
            run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(WML_NS + 'type', 'page')
            run.append(br)
            page_break_para.append(run)
            
            body.append(page_break_para)
            
            logger.debug("双面打印：已在末尾添加分页")


# 导入其他模块
from .style_mapper import StyleMapper
from .numbering import NumberingManager
from .toc_generator import TOCGenerator
from .cross_reference import CrossReferenceManager
