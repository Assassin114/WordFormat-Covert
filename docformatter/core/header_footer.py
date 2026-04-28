"""
页眉页脚管理器
支持三类页面的不同页眉页脚设置：
1. 封面页、签署页、修改记录页：无页码
2. 目录页：大写罗马数字页码，从I开始
3. 正文部分：大写罗马数字页码，从I开始（新编号）
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from ..utils.logger import get_logger

logger = get_logger()

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class HeaderFooterManager:
    """
    页眉页脚管理器
    """
    
    TYPE_COVER = "cover"
    TYPE_TOC = "toc"
    TYPE_BODY = "body"
    
    def __init__(self, template=None):
        self.template = template
    
    def apply(self, doc: Document):
        """应用页眉页脚设置"""
        section_types = self._analyze_document_structure(doc)
        
        logger.info(f"文档结构分析: {len(section_types)} 个节")
        for i, sect_type in enumerate(section_types):
            logger.debug(f"  节{i}: {sect_type}")
        
        for i, (section, sect_type) in enumerate(zip(doc.sections, section_types)):
            self._apply_header_footer_for_section(section, sect_type, i)
    
    def _analyze_document_structure(self, doc) -> list:
        """分析文档结构，返回每节的类型列表"""
        section_types = []
        current_type = self.TYPE_COVER
        section_starts = self._get_section_starts(doc)
        section_count = 0
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if i in section_starts and section_count > 0:
                section_types.append(current_type)
                current_type = self.TYPE_COVER
            
            section_count = i // max(1, len(doc.paragraphs) // max(1, len(section_starts) or 1))
            
            if self._is_toc_heading(text):
                current_type = self.TYPE_TOC
            
            if self._is_body_start(text):
                current_type = self.TYPE_BODY
        
        section_types.append(current_type)
        
        if len(section_types) == 1:
            section_types = self._refine_single_section(doc, section_types[0])
        
        return section_types
    
    def _get_section_starts(self, doc) -> set:
        """获取所有节的起始段落索引"""
        starts = set()
        for i, para in enumerate(doc.paragraphs):
            pPr = para._element.find(WML_NS + 'pPr')
            if pPr is not None:
                sectPr = pPr.find(WML_NS + 'sectPr')
                if sectPr is not None:
                    starts.add(i)
        return starts
    
    def _is_toc_heading(self, text: str) -> bool:
        """判断是否为目录标题"""
        toc_markers = ['目录', 'CONTENTS', 'Table of Contents', '目录页']
        return text in toc_markers or '目  录' in text
    
    def _is_body_start(self, text: str) -> bool:
        """判断是否为正文开始"""
        body_markers = ['第一章', '1.', '一、', '摘要', 'Abstract', '第1章', '引言', 'Introduction']
        return any(text.startswith(m) for m in body_markers)
    
    def _refine_single_section(self, doc, default_type: str) -> list:
        """细化单节文档的类型判断"""
        has_toc = any(self._is_toc_heading(p.text.strip()) for p in doc.paragraphs)
        has_body = any(self._is_body_start(p.text.strip()) for p in doc.paragraphs)
        
        if has_toc and has_body:
            return [self.TYPE_COVER, self.TYPE_TOC, self.TYPE_BODY]
        elif has_toc:
            return [self.TYPE_COVER, self.TYPE_TOC]
        elif has_body:
            return [self.TYPE_COVER, self.TYPE_BODY]
        else:
            return [self.TYPE_COVER]
    
    def _apply_header_footer_for_section(self, section, sect_type: str, section_index: int):
        """为指定节应用页眉页脚"""
        header = section.header
        footer = section.footer
        
        for para in header.paragraphs:
            para.clear()
        for para in footer.paragraphs:
            para.clear()
        
        if sect_type == self.TYPE_COVER:
            self._set_header_text(header, self._get_cover_header_text())
        elif sect_type == self.TYPE_TOC:
            self._set_header_text(header, self._get_toc_header_text())
            self._set_footer_with_page_number(footer, use_roman=True, start_at=1)
        elif sect_type == self.TYPE_BODY:
            self._set_header_text(header, self._get_body_header_text())
            self._set_footer_with_page_number(footer, use_roman=True, start_at=1)
    
    def _get_cover_header_text(self) -> str:
        if self.template and hasattr(self.template, 'header_footer'):
            return getattr(self.template.header_footer, 'cover_header_text', '文档标题')
        return '文档标题'
    
    def _get_toc_header_text(self) -> str:
        if self.template and hasattr(self.template, 'header_footer'):
            return getattr(self.template.header_footer, 'toc_header_text', '目录')
        return '目录'
    
    def _get_body_header_text(self) -> str:
        if self.template and hasattr(self.template, 'header_footer'):
            return getattr(self.template.header_footer, 'body_header_text', '正文')
        return '正文'
    
    def _set_header_text(self, header, text: str):
        """设置页眉文本"""
        if header.paragraphs:
            para = header.paragraphs[0]
        else:
            para = header.add_paragraph()
        para.clear()
        run = para.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        para.alignment = 1
    
    def _set_footer_with_page_number(self, footer, use_roman: bool = True, start_at: int = 1):
        """设置带页码的页脚"""
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.clear()
        para.alignment = 1
        
        run = para.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(WML_NS + 'fldCharType', 'begin')
        run._r.append(fld_char_begin)
        
        run2 = para.add_run()
        instr_text = OxmlElement('w:instrText')
        instr_text.set(WML_NS + 'xml:space', 'preserve')
        instr_text.text = ' PAGE \\* roman '
        run2._r.append(instr_text)
        
        run3 = para.add_run()
        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(WML_NS + 'fldCharType', 'separate')
        run3._r.append(fld_char_sep)
        
        run4 = para.add_run()
        t = OxmlElement('w:t')
        t.text = str(start_at)
        run4._r.append(t)
        
        run5 = para.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(WML_NS + 'fldCharType', 'end')
        run5._r.append(fld_char_end)
